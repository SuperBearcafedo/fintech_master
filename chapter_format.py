# -*- coding: utf-8 -*-
import os
import sys
import traceback
import base_import
import docx
from docx import Document
from docx.shared import Length, Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_UNDERLINE
from datetime import datetime


class Chapter:  # 类采用Pascal命名风格
    """
    目前缺乏对类使用的说明：后续进行补充
    方法可以在该类下继续添加，逐渐丰富
    """
    def __init__(self, create_path, add_file_name, title_name, chapter_name):
        self.document_object = Document()  # 创建文件对象
        self.create_path = create_path
        self.add_file_name = add_file_name
        self.tn = title_name
        self.cn = chapter_name

    def createTitle(self):  # 小驼峰命名法：局部使用
        """
        入参：self
        :return: none
        """
        self.existChapterFile()
        # 创建文件基本对象信息
        run_heading = self.document_object.add_heading("", level=1)
        run = run_heading.add_run(self.tn)  # 以追加段落格式方便后续修改
        run.font.underline = False  # 不需要添加下划线
        run_heading.paragraph_format.space_before = Pt(0)  # 设置段前
        run_heading.paragraph_format.space_after = Pt(0)  # 设置段后
        run_heading.paragraph_format.line_spacing = 1.5  # 设置行间距
        run.font.size = Pt(20)  # 设置1级标题文字的大小为“小四” 为12磅
        run.font.name = "Times New Roman"  # 设置英文字体
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u"宋体")  # 设置中文字体
        run.font.color.rgb = RGBColor(0, 0, 0)  # 设置颜色

        # 创建内容
        content = self.document_object.add_paragraph()
        content.paragraph_format.space_before = Pt(10)  # 设置段前 0 磅
        content.paragraph_format.space_after = Pt(10)  # 设置段后 0 磅
        content.paragraph_format.line_spacing = 1.5  # 设置行间距为 1.5倍
        content.paragraph_format.first_line_indent = Inches(0)  # 段落首行缩进为 0.5英寸
        content.paragraph_format.first_line_indent = Inches(0)  # 相当于小四两个字符的缩进
        content_run = content.add_run(self.cn)
        content_run.bold = False  # 加粗
        content_run.font.size = Pt(16)
        content_run.font.name = "Times New Roman"  # 设置英文字体
        content_run._element.rPr.rFonts.set(qn('w:eastAsia'), u"宋体")  # 设置中文字体
        content_run.font.color.rgb = RGBColor(0, 0, 0)  # 设置颜色

        # 创建更新时间，添加时间
        header_section = self.document_object.sections[0]
        header_pos = header_section.header
        header_para = header_pos.paragraphs[0]
        current_date_and_time = datetime.now().strftime("%y-%m-%d %H:%M:%S")
        header_para.text = "该章节编写于:" + str(current_date_and_time)

        self.saveChapter()

        return

    def saveChapter(self):
        save_path = os.path.join(self.create_path, self.add_file_name)
        self.document_object.save(save_path)  # 创建文件保存地址

    def existChapterFile(self):
        if not os.path.exists(self.create_path):
            print("章节创建文件夹路径不存在，创建目录: {}".format(self.create_path))
            os.mkdir(create_path)


if __name__ == "__main__":
    create_path = os.path.join(os.getcwd(), "project_file")
    add_file_name = "summarybook.docx"
    title_name = "第一章/First Chapter"
    chapter_name = "深度学习Python从入门到实践"

    first_chapter = Chapter(create_path, add_file_name, title_name, chapter_name)
    first_chapter.createTitle()
    first_chapter.createTable()
